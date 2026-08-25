from io import BytesIO
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    ListFlowable,
    ListItem,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
)


# =========================================================
# Shared helpers
# =========================================================

def _clean_text(value) -> str:
    if value is None:
        return ""

    text = str(value).strip()

    replacements = {
        "\u2013": "-",
        "\u2014": "-",
        "\u2018": "'",
        "\u2019": "'",
        "\u201c": '"',
        "\u201d": '"',
        "\u00a0": " ",
    }

    for old, new in replacements.items():
        text = text.replace(old, new)

    return text


# =========================================================
# DOCX
# =========================================================

def _configure_docx(document: Document) -> None:
    section = document.sections[0]

    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.5)

    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.0


def _add_docx_section_heading(
    document: Document,
    title: str,
) -> None:
    paragraph = document.add_paragraph()

    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(3)

    run = paragraph.add_run(
        _clean_text(title).upper()
    )

    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(10.5)


def _add_docx_bullet(
    document: Document,
    text: str,
) -> None:
    paragraph = document.add_paragraph(
        style="List Bullet"
    )

    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.first_line_indent = Inches(-0.12)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.0

    run = paragraph.add_run(
        _clean_text(text)
    )

    run.font.name = "Arial"
    run.font.size = Pt(9.2)


def render_tailored_cv_docx(
    candidate_name: str,
    tailored_cv: dict,
) -> bytes:
    document = Document()

    _configure_docx(document)

    # -----------------------------------------------------
    # Header
    # -----------------------------------------------------

    name_paragraph = document.add_paragraph()
    name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_paragraph.paragraph_format.space_after = Pt(2)

    run = name_paragraph.add_run(
        _clean_text(candidate_name).upper()
    )

    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(17)

    headline = _clean_text(
        tailored_cv.get("headline", "")
    )

    if headline:
        headline_paragraph = document.add_paragraph()
        headline_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        headline_paragraph.paragraph_format.space_after = Pt(7)

        run = headline_paragraph.add_run(headline)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(10.2)

    # -----------------------------------------------------
    # Summary
    # -----------------------------------------------------

    summary = _clean_text(
        tailored_cv.get(
            "professional_summary",
            "",
        )
    )

    if summary:
        _add_docx_section_heading(
            document,
            "Professional Summary",
        )

        paragraph = document.add_paragraph(summary)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.0

    # -----------------------------------------------------
    # Skills
    # -----------------------------------------------------

    skills = [
        _clean_text(item)
        for item in tailored_cv.get(
            "key_skills",
            [],
        )
        if _clean_text(item)
    ]

    if skills:
        _add_docx_section_heading(
            document,
            "Core Skills",
        )

        paragraph = document.add_paragraph(
            " | ".join(skills)
        )

        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.0

    # -----------------------------------------------------
    # Experience
    # -----------------------------------------------------

    experiences = tailored_cv.get(
        "experiences",
        [],
    )

    if experiences:
        _add_docx_section_heading(
            document,
            "Professional Experience",
        )

        for experience in experiences:
            role = _clean_text(
                experience.get("role", "")
            )

            company = _clean_text(
                experience.get("company", "")
            )

            heading = " | ".join(
                value
                for value in [
                    company,
                    role,
                ]
                if value
            )

            if heading:
                paragraph = document.add_paragraph()

                paragraph.paragraph_format.space_before = Pt(4)
                paragraph.paragraph_format.space_after = Pt(2)

                run = paragraph.add_run(heading)
                run.bold = True
                run.font.name = "Arial"
                run.font.size = Pt(9.8)

            for bullet in experience.get(
                "tailored_bullets",
                [],
            ):
                if _clean_text(bullet):
                    _add_docx_bullet(
                        document,
                        bullet,
                    )

    # -----------------------------------------------------
    # Additional information
    # -----------------------------------------------------

    additional = [
        _clean_text(item)
        for item in tailored_cv.get(
            "additional_relevant_information",
            [],
        )
        if _clean_text(item)
    ]

    if additional:
        _add_docx_section_heading(
            document,
            "Additional Information",
        )

        for item in additional:
            _add_docx_bullet(
                document,
                item,
            )

    buffer = BytesIO()
    document.save(buffer)

    data = buffer.getvalue()

    if not data.startswith(b"PK"):
        raise ValueError(
            "Generated DOCX is invalid."
        )

    return data


# =========================================================
# PDF
# =========================================================

def render_tailored_cv_pdf(
    candidate_name: str,
    tailored_cv: dict,
) -> bytes:
    buffer = BytesIO()

    document = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=15 * mm,
        leftMargin=15 * mm,
        topMargin=13 * mm,
        bottomMargin=13 * mm,
    )

    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "CVTitle",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=16,
        leading=18,
        alignment=TA_CENTER,
        spaceAfter=2,
    )

    headline_style = ParagraphStyle(
        "CVHeadline",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=9.5,
        leading=12,
        alignment=TA_CENTER,
        spaceAfter=8,
    )

    section_style = ParagraphStyle(
        "CVSection",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=10,
        leading=12,
        spaceBefore=7,
        spaceAfter=3,
    )

    body_style = ParagraphStyle(
        "CVBody",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=8.8,
        leading=11.2,
        spaceAfter=3,
    )

    role_style = ParagraphStyle(
        "CVRole",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=9.2,
        leading=11.5,
        spaceBefore=4,
        spaceAfter=2,
    )

    story = []

    def pdf_text(value) -> str:
        return escape(
            _clean_text(value)
        )

    # -----------------------------------------------------
    # Header
    # -----------------------------------------------------

    story.append(
        Paragraph(
            pdf_text(candidate_name).upper(),
            title_style,
        )
    )

    headline = _clean_text(
        tailored_cv.get("headline", "")
    )

    if headline:
        story.append(
            Paragraph(
                pdf_text(headline),
                headline_style,
            )
        )

    # -----------------------------------------------------
    # Summary
    # -----------------------------------------------------

    summary = _clean_text(
        tailored_cv.get(
            "professional_summary",
            "",
        )
    )

    if summary:
        story.append(
            Paragraph(
                "PROFESSIONAL SUMMARY",
                section_style,
            )
        )

        story.append(
            Paragraph(
                pdf_text(summary),
                body_style,
            )
        )

    # -----------------------------------------------------
    # Skills
    # -----------------------------------------------------

    skills = [
        _clean_text(item)
        for item in tailored_cv.get(
            "key_skills",
            [],
        )
        if _clean_text(item)
    ]

    if skills:
        story.append(
            Paragraph(
                "CORE SKILLS",
                section_style,
            )
        )

        story.append(
            Paragraph(
                pdf_text(
                    " | ".join(skills)
                ),
                body_style,
            )
        )

    # -----------------------------------------------------
    # Experience
    # -----------------------------------------------------

    experiences = tailored_cv.get(
        "experiences",
        [],
    )

    if experiences:
        story.append(
            Paragraph(
                "PROFESSIONAL EXPERIENCE",
                section_style,
            )
        )

        for experience in experiences:
            role = _clean_text(
                experience.get("role", "")
            )

            company = _clean_text(
                experience.get("company", "")
            )

            heading = " | ".join(
                value
                for value in [
                    company,
                    role,
                ]
                if value
            )

            if heading:
                story.append(
                    Paragraph(
                        pdf_text(heading),
                        role_style,
                    )
                )

            bullets = [
                _clean_text(item)
                for item in experience.get(
                    "tailored_bullets",
                    [],
                )
                if _clean_text(item)
            ]

            if bullets:
                story.append(
                    ListFlowable(
                        [
                            ListItem(
                                Paragraph(
                                    pdf_text(bullet),
                                    body_style,
                                ),
                                leftIndent=8,
                            )
                            for bullet in bullets
                        ],
                        bulletType="bullet",
                        leftIndent=13,
                        bulletFontName="Helvetica",
                        bulletFontSize=5.5,
                        spaceAfter=0.5,
                    )
                )

    # -----------------------------------------------------
    # Additional
    # -----------------------------------------------------

    additional = [
        _clean_text(item)
        for item in tailored_cv.get(
            "additional_relevant_information",
            [],
        )
        if _clean_text(item)
    ]

    if additional:
        story.append(
            Paragraph(
                "ADDITIONAL INFORMATION",
                section_style,
            )
        )

        story.append(
            ListFlowable(
                [
                    ListItem(
                        Paragraph(
                            pdf_text(item),
                            body_style,
                        ),
                        leftIndent=8,
                    )
                    for item in additional
                ],
                bulletType="bullet",
                leftIndent=13,
                bulletFontSize=6,
            )
        )

    document.build(story)

    data = buffer.getvalue()

    if not data.startswith(b"%PDF-"):
        raise ValueError(
            "Generated PDF is invalid."
        )

    return data
